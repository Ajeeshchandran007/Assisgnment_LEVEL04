"""
LLM-LangChain-MongoDB Toolchain - LOCAL MongoDB Compatible
Works with local MongoDB instance without requiring Atlas
Uses in-memory vector search with MongoDB for document storage
"""

import os
import numpy as np
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from pymongo import MongoClient
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.vectorstores import VectorStore
from langchain_core.embeddings import Embeddings
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

class LocalMongoDBVectorStore(VectorStore):
    """
    Custom Vector Store that works with local MongoDB
    Stores embeddings in MongoDB and performs similarity search in-memory
    """
    
    def __init__(self, collection, embedding: Embeddings, text_key: str = "text",embedding_key: str = "embedding"):
        self.collection = collection
        self.embedding_function = embedding
        self.text_key = text_key
        self.embedding_key = embedding_key
    
    def add_texts(self, texts: List[str], metadatas: Optional[List[dict]] = None, **kwargs) -> List[str]:
        """Add texts to the vector store"""
        if metadatas is None:
            metadatas = [{} for _ in texts]
        
        # Generate embeddings
        embeddings = self.embedding_function.embed_documents(texts)
        
        # Prepare documents
        documents = []
        for text, embedding, metadata in zip(texts, embeddings, metadatas):
            doc = {
                self.text_key: text,
                self.embedding_key: embedding,
                "metadata": metadata,
                "created_at": datetime.now().isoformat()
            }
            documents.append(doc)
        
        # Insert into MongoDB
        result = self.collection.insert_many(documents)
        return [str(id) for id in result.inserted_ids]
    
    def similarity_search(self, query: str, k: int = 4, **kwargs) -> List[Document]:
        """Perform similarity search"""
        docs_and_scores = self.similarity_search_with_score(query, k=k)
        return [doc for doc, _ in docs_and_scores]
    
    def similarity_search_with_score(self, query: str, k: int = 4,**kwargs) -> List[Tuple[Document, float]]:
        """Perform similarity search with scores"""
        # Generate query embedding
        query_embedding = self.embedding_function.embed_query(query)
        
        # Retrieve all documents from MongoDB
        all_docs = list(self.collection.find())
        
        if not all_docs:
            return []
        
        # Calculate cosine similarity
        similarities = []
        for doc in all_docs:
            if self.embedding_key in doc:
                doc_embedding = doc[self.embedding_key]
                similarity = self._cosine_similarity(query_embedding, doc_embedding)
                similarities.append((doc, similarity))
        
        # Sort by similarity (highest first)
        similarities.sort(key=lambda x: x[1], reverse=True)
        
        # Return top k results
        results = []
        for doc, score in similarities[:k]:
            langchain_doc = Document(
                page_content=doc[self.text_key],
                metadata=doc.get("metadata", {})
            )
            results.append((langchain_doc, score))
        
        return results
    
    @staticmethod
    def _cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors"""
        vec1 = np.array(vec1)
        vec2 = np.array(vec2)
        
        dot_product = np.dot(vec1, vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        
        if norm1 == 0 or norm2 == 0:
            return 0.0
        
        return float(dot_product / (norm1 * norm2))
    
    @classmethod
    def from_texts(cls,texts: List[str], embedding: Embeddings, metadatas: Optional[List[dict]] = None,collection=None, **kwargs):
        """Create vector store from texts - required by VectorStore base class"""
        if collection is None:
            raise ValueError("collection parameter is required")
        
        vector_store = cls(collection=collection, embedding=embedding)
        vector_store.add_texts(texts, metadatas)
        
        return vector_store
    
    @classmethod
    def from_documents(cls, documents: List[Document], embedding: Embeddings, collection, **kwargs):
        """Create vector store from documents"""
        texts = [doc.page_content for doc in documents]
        metadatas = [doc.metadata for doc in documents]
        
        return cls.from_texts(
            texts=texts,
            embedding=embedding,
            metadatas=metadatas,
            collection=collection,
            **kwargs
        )
    
    def _similarity_search_with_relevance_scores(self, query: str, k: int = 4, **kwargs) -> List[Tuple[Document, float]]:
        """Required by VectorStore base class"""
        return self.similarity_search_with_score(query, k=k, **kwargs)


class LLMMongoDBToolchain:
    """
    Toolchain for RAG-based querying with LOCAL MongoDB and LangChain 1.0+
    Works without MongoDB Atlas - uses local MongoDB for storage
    """
    
    def __init__(self, openai_api_key: str, mongo_uri: str = "mongodb://localhost:27017/", database_name: str = "llm_database", collection_name: str = "documents"):
        """
        Initialize the toolchain with LangChain 1.0+ components
        
        Args:
            openai_api_key: OpenAI API key
            mongo_uri: MongoDB connection string
            database_name: Name of the MongoDB database
            collection_name: Name of the collection for documents
        """
        # Set OpenAI API key
        os.environ["OPENAI_API_KEY"] = openai_api_key
        
        # Initialize MongoDB client
        self.client = MongoClient(mongo_uri)
        self.db = self.client[database_name]
        self.collection = self.db[collection_name]
        
        # Initialize embeddings with latest model
        self.embeddings = OpenAIEmbeddings(
            model="text-embedding-3-small",
            dimensions=1536
        )
        
        # Initialize LLM
        self.llm = ChatOpenAI(
            model="gpt-4o-mini",
            temperature=0,
            max_tokens=500
        )
        
        # Initialize vector store
        self.vector_store: Optional[LocalMongoDBVectorStore] = None
        
        print(f"✓ Connected to LOCAL MongoDB at {mongo_uri}")
        print(f"✓ Using database: {self.db.name}")
        print(f"✓ Using collection: {self.collection.name}")
        print(f"✓ LangChain version: 1.0+ compatible")
        print(f"✓ Vector search: Local (in-memory)")
    
    def upload_documents(self, documents: List[Dict[str, str]]):
        """
        Upload documents to MongoDB with embeddings
        
        Args:
            documents: List of dicts with 'content' and optional 'metadata'
        """
        print(f"\n📤 Uploading {len(documents)} documents...")
        
        # Convert to LangChain Document objects
        docs = []
        for doc in documents:
            metadata = doc.get('metadata', {})
            metadata['upload_time'] = datetime.now().isoformat()
            metadata['source'] = doc.get('source', 'user_upload')
            
            docs.append(Document(
                page_content=doc['content'],
                metadata=metadata
            ))
        
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""],
            keep_separator=False
        )
        split_docs = text_splitter.split_documents(docs)
        
        print(f"✓ Split into {len(split_docs)} chunks")
        
        # Create vector store and add documents
        self.vector_store = LocalMongoDBVectorStore.from_documents(
            documents=split_docs,
            embedding=self.embeddings,
            collection=self.collection
        )
        
        print(f"✓ Uploaded and embedded {len(split_docs)} document chunks")
        return len(split_docs)
    
    def _format_docs(self, docs: List[Document]) -> str:
        """Format documents for context"""
        return "\n\n".join([
            f"Document {i+1}:\n{doc.page_content}" 
            for i, doc in enumerate(docs)
        ])
    
    def query_with_llm(self, query: str, k: int = 3, return_source_docs: bool = True) -> Dict[str, Any]:
        """
        Query the database using LLM with RAG - LangChain 1.0+ LCEL pattern
        
        Args:
            query: Query string
            k: Number of relevant documents to retrieve
            return_source_docs: Whether to return source documents
            
        Returns:
            Dict with answer and optionally source documents
        """
        if self.vector_store is None:
            # Initialize vector store from existing collection
            self.vector_store = LocalMongoDBVectorStore(
                collection=self.collection,
                embedding=self.embeddings
            )
        
        # Create retriever
        retriever = self.vector_store.as_retriever(
            search_kwargs={"k": k}
        )
        
        # Create RAG prompt
        template = """You are a helpful AI assistant. Use the following pieces of context to answer the question at the end.
If you don't know the answer based on the context, just say that you don't know, don't try to make up an answer.

Context:
{context}

Question: {question}

Answer:"""
        
        prompt = ChatPromptTemplate.from_template(template)
        
        # Build RAG chain using LCEL
        rag_chain = (
            {
                "context": retriever | self._format_docs,
                "question": RunnablePassthrough()
            }
            | prompt
            | self.llm
            | StrOutputParser()
        )
        
        # Execute query
        print(f"\n🔍 Querying: '{query}'")
        
        # Get answer
        answer = rag_chain.invoke(query)
        
        result = {"answer": answer}
        
        # Optionally get source documents
        if return_source_docs:
            source_docs = retriever.invoke(query)
            result["source_documents"] = source_docs
            result["num_sources"] = len(source_docs)
        
        return result
    
    def query_with_streaming(self, query: str, k: int = 3):
        """
        Query with streaming response - LangChain 1.0+ streaming pattern
        
        Args:
            query: Query string
            k: Number of relevant documents to retrieve
        """
        if self.vector_store is None:
            self.vector_store = LocalMongoDBVectorStore(
                collection=self.collection,
                embedding=self.embeddings
            )
        
        retriever = self.vector_store.as_retriever(
            search_kwargs={"k": k}
        )
        
        template = """Use the following context to answer the question.

Context:
{context}

Question: {question}

Answer:"""
        
        prompt = ChatPromptTemplate.from_template(template)
        
        # Build streaming chain
        rag_chain = (
            {
                "context": retriever | self._format_docs,
                "question": RunnablePassthrough()
            }
            | prompt
            | self.llm
            | StrOutputParser()
        )
        
        print(f"\n🔍 Streaming query: '{query}'")
        print("\n📝 Response: ", end="", flush=True)
        
        # Stream the response
        for chunk in rag_chain.stream(query):
            print(chunk, end="", flush=True)
        
        print("\n")
    
    def display_results(self, result: Dict[str, Any]):
        """Display query results in a formatted way"""
        print("\n" + "="*80)
        print("ANSWER:")
        print("="*80)
        print(result["answer"])
        
        if "source_documents" in result:
            print("\n" + "="*80)
            print(f"SOURCE DOCUMENTS ({result.get('num_sources', 0)} retrieved):")
            print("="*80)
            for i, doc in enumerate(result["source_documents"], 1):
                print(f"\n[Source {i}]")
                print(f"Content: {doc.page_content[:200]}...")
                if doc.metadata:
                    print(f"Metadata: {doc.metadata}")
    
    def similarity_search(self, query: str, k: int = 3) -> List[Tuple[Document, float]]:
        """
        Perform similarity search without LLM (with scores)
        
        Args:
            query: Query string
            k: Number of results to return
            
        Returns:
            List of (document, similarity_score) tuples
        """
        if self.vector_store is None:
            self.vector_store = LocalMongoDBVectorStore(
                collection=self.collection,
                embedding=self.embeddings
            )
        
        return self.vector_store.similarity_search_with_score(query, k=k)
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the MongoDB collection"""
        return {
            "total_documents": self.collection.count_documents({}),
            "database": self.db.name,
            "collection": self.collection.name,
            "indexes": [idx['name'] for idx in self.collection.list_indexes()],
            "vector_store_type": "Local (in-memory search)"
        }
    
    def clear_collection(self):
        """Clear all documents from the collection"""
        result = self.collection.delete_many({})
        print(f"✓ Deleted {result.deleted_count} documents")
        self.vector_store = None
    
    def close(self):
        """Close MongoDB connection"""
        self.client.close()
        print("✓ MongoDB connection closed")
    
    def load_documents_from_folder(self, folder_path: str = "./documents") -> List[Dict[str, str]]:
        """
        Load all documents from a folder including PDF, DOCX, and text files
        
        Args:
            folder_path: Path to the folder containing documents
            
        Returns:
            List of document dictionaries with content and metadata
        """
        import glob
        
        documents = []
        supported_extensions = ['.txt', '.md', '.json', '.py', '.js', '.html', '.css', '.pdf', '.docx', '.doc']
        
        if not os.path.exists(folder_path):
            print(f"❌ Folder '{folder_path}' does not exist!")
            return documents
        
        # Find all supported files
        file_patterns = [os.path.join(folder_path, f"*{ext}") for ext in supported_extensions]
        files = []
        for pattern in file_patterns:
            files.extend(glob.glob(pattern))
        
        if not files:
            print(f"⚠️ No supported files found in '{folder_path}'")
            print(f"Supported formats: {', '.join(supported_extensions)}")
            return documents
        
        print(f"\n📂 Loading documents from '{folder_path}'...")
        print(f"Found {len(files)} files")
        
        for file_path in files:
            try:
                filename = os.path.basename(file_path)
                file_extension = os.path.splitext(filename)[1].lower()
                content = None
                
                # Handle PDF files
                if file_extension == '.pdf':
                    content = self._extract_pdf_content(file_path)
                
                # Handle DOCX files
                elif file_extension == '.docx':
                    content = self._extract_docx_content(file_path)
                
                # Handle DOC files (older Word format)
                elif file_extension == '.doc':
                    content = self._extract_doc_content(file_path)
                
                # Handle text-based files
                else:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                
                # Skip empty files
                if not content or not content.strip():
                    print(f"⚠️ Skipping empty file: {filename}")
                    continue
                
                documents.append({
                    "content": content,
                    "metadata": {
                        "source": file_path,
                        "filename": filename,
                        "file_type": file_extension,
                        "file_size": len(content),
                        "loaded_at": datetime.now().isoformat()
                    }
                })
                
                print(f"✓ Loaded: {filename} ({len(content)} characters)")
                
            except Exception as e:
                print(f"❌ Error loading {file_path}: {str(e)}")
                continue
        
        print(f"\n✅ Successfully loaded {len(documents)} documents")
        return documents
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """
        Extract text content from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        try:
            import PyPDF2
            
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{text}")
            
            return "\n\n".join(text_content)
        
        except ImportError:
            raise ImportError(
                "PyPDF2 is required to read PDF files. "
                "Install it with: pip install PyPDF2"
            )
        except Exception as e:
            raise Exception(f"Error extracting PDF content: {str(e)}")
    
    def _extract_docx_content(self, file_path: str) -> str:
        """
        Extract text content from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            return "\n\n".join(text_content)
        
        except ImportError:
            raise ImportError(
                "python-docx is required to read DOCX files. "
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            raise Exception(f"Error extracting DOCX content: {str(e)}")
    
    def _extract_doc_content(self, file_path: str) -> str:
        """
        Extract text content from DOC file (older Word format)
        
        Args:
            file_path: Path to DOC file
            
        Returns:
            Extracted text content
        """
        try:
            import textract
            
            # textract can handle .doc files
            text = textract.process(file_path).decode('utf-8')
            return text
        
        except ImportError:
            raise ImportError(
                "textract is required to read DOC files. "
                "Install it with: pip install textract\n"
                "Note: textract may require additional system dependencies."
            )
        except Exception as e:
            raise Exception(f"Error extracting DOC content: {str(e)}")


# Example usage demonstrating LangChain 1.0+ features
def main():
    """Example demonstration of the LangChain 1.0+ toolchain with LOCAL MongoDB"""
    
    # Load OpenAI API key from environment
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
    
    if not OPENAI_API_KEY:
        print("❌ ERROR: OpenAI API key not found!")
        print("\n📝 Please create a .env file in the same directory with:")
        print("   OPENAI_API_KEY=sk-your-actual-api-key-here")
        print("\n📌 Steps to set up:")
        print("   1. Create a file named '.env' (with the dot)")
        print("   2. Add the line: OPENAI_API_KEY=sk-your-key-here")
        print("   3. Save the file")
        print("   4. Run this script again")
        print("\n🔑 Get your API key from: https://platform.openai.com/api-keys")
        return
    
    MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017/")
    
    print("✅ OpenAI API key loaded from .env file")
    
    toolchain = LLMMongoDBToolchain(
        openai_api_key=OPENAI_API_KEY,
        mongo_uri=MONGO_URI,
        database_name="llm_database_local",
        collection_name="documents_local"
    )
    
    # Load documents from ./documents folder
    print("\n" + "="*80)
    print("LOADING DOCUMENTS FROM FOLDER")
    print("="*80)
    
    documents = toolchain.load_documents_from_folder("./documents")
    
    if not documents:
        print("\n⚠️ No documents found! Using sample documents instead...")
        # Fallback to sample documents if folder is empty
        documents = [
            {
                "content": """
                Large Language Models (LLMs) are neural networks with billions of parameters 
                trained on vast amounts of text data. They can perform various natural language 
                tasks like translation, summarization, and question-answering. Popular examples 
                include GPT-4, Claude, and Gemini. These models use transformer architecture 
                and attention mechanisms to understand context and generate human-like text.
                LLMs have revolutionized AI applications across industries.
                """,
                "metadata": {"topic": "LLMs", "category": "AI", "year": 2024}
            }
        ]
    
    # Upload documents
    print("\n" + "="*80)
    print("UPLOADING DOCUMENTS")
    print("="*80)
    toolchain.upload_documents(documents)
    
    # Show collection stats
    stats = toolchain.get_collection_stats()
    print(f"\n📊 Collection Stats:")
    for key, value in stats.items():
        print(f"  {key}: {value}")
    
    # Interactive query loop
    print("\n" + "="*80)
    print("INTERACTIVE QUERY MODE")
    print("="*80)
    print("\n💡 You can now ask questions about your documents!")
    print("Commands:")
    print("  - Type your question to get an answer")
    print("  - Type 'exit' or 'quit' to end the session")
    print("  - Type 'clear' to clear the collection")
    print("  - Type 'stats' to see collection statistics")
    print("="*80)
    
    while True:
        try:
            # Get user input
            print("\n")
            user_query = input("🤔 Your question: ").strip()
            
            # Check for exit commands
            if user_query.lower() in ['exit', 'quit', 'q']:
                print("\n👋 Thank you for using the RAG system! Goodbye!")
                break
            
            # Check for empty input
            if not user_query:
                print("⚠️ Please enter a question.")
                continue
            
            # Handle special commands
            if user_query.lower() == 'clear':
                confirm = input("⚠️ Are you sure you want to clear all documents? (yes/no): ")
                if confirm.lower() == 'yes':
                    toolchain.clear_collection()
                    print("✅ Collection cleared. You can upload new documents.")
                continue
            
            if user_query.lower() == 'stats':
                stats = toolchain.get_collection_stats()
                print("\n📊 Collection Stats:")
                for key, value in stats.items():
                    print(f"  {key}: {value}")
                continue
            
            # Process the query (without source documents)
            result = toolchain.query_with_llm(user_query, k=3, return_source_docs=False)
            
            # Display only the answer
            print("\n" + "="*80)
            print("ANSWER:")
            print("="*80)
            print(result["answer"])
            print("="*80)
            
        except KeyboardInterrupt:
            print("\n\n👋 Session interrupted. Goodbye!")
            break
        except Exception as e:
            print(f"\n❌ Error processing query: {str(e)}")
            print("Please try again or type 'exit' to quit.")
    
    # Close connection
    print("\n" + "="*80)
    toolchain.close()
    
    print("\n✅ Demo completed successfully!")
    print("\n💡 TIP: This works with LOCAL MongoDB - no Atlas required!")


if __name__ == "__main__":
    main()